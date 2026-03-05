"""
Document text extraction for OpenScientist.

Extracts text from binary document formats (PDF, DOCX, XLSX) that cannot
be read directly by Claude's Read tool without returning garbled content.
"""

import logging
from pathlib import Path
from typing import Any

import fitz  # type: ignore[import-untyped]  # PyMuPDF
import openpyxl
from docx import Document

logger = logging.getLogger(__name__)

# File extensions that require special extraction
BINARY_DOCUMENT_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".doc",
    ".xlsx",
    ".xls",
    ".pptx",
    ".ppt",
}

# Extensions that can be read as text directly
TEXT_EXTENSIONS = {
    ".csv",
    ".tsv",
    ".txt",
    ".json",
    ".jsonl",
    ".md",
    ".markdown",
    ".xml",
    ".html",
    ".htm",
    ".yaml",
    ".yml",
    ".py",
    ".r",
    ".sh",
}


def is_binary_document(file_path: Path) -> bool:
    """Check if a file is a binary document that needs special extraction."""
    return file_path.suffix.lower() in BINARY_DOCUMENT_EXTENSIONS


def is_text_file(file_path: Path) -> bool:
    """Check if a file can be read as plain text."""
    return file_path.suffix.lower() in TEXT_EXTENSIONS


def extract_text_from_pdf(
    file_path: Path, max_pages: int | None = None
) -> tuple[str, dict[str, Any]]:
    """
    Extract text from a PDF file using PyMuPDF.

    Args:
        file_path: Path to PDF file
        max_pages: Maximum number of pages to extract (None for all)

    Returns:
        Tuple of (extracted_text, metadata_dict)
    """
    doc = fitz.open(str(file_path))
    metadata = {
        "format": "pdf",
        "pages": len(doc),
        "title": doc.metadata.get("title") or "",
        "author": doc.metadata.get("author") or "",
    }

    text_parts: list[str] = []
    pages_to_extract = min(len(doc), max_pages) if max_pages else len(doc)

    for page_num in range(pages_to_extract):
        page = doc[page_num]
        text = page.get_text()
        if text.strip():
            text_parts.append(f"--- Page {page_num + 1} ---\n{text}")

    total_pages = len(doc)
    doc.close()

    extracted_text = "\n\n".join(text_parts)

    # Add truncation notice if we didn't extract all pages
    if max_pages and total_pages > max_pages:
        extracted_text += f"\n\n[... Truncated. Showing {max_pages} of {total_pages} pages ...]"

    return extracted_text, metadata


def extract_text_from_docx(file_path: Path) -> tuple[str, dict[str, Any]]:
    """
    Extract text from a Word document using python-docx.

    Args:
        file_path: Path to DOCX file

    Returns:
        Tuple of (extracted_text, metadata_dict)
    """
    doc = Document(str(file_path))
    metadata = {
        "format": "docx",
        "paragraphs": len(doc.paragraphs),
    }

    # Extract core properties if available
    try:
        if doc.core_properties:
            metadata["title"] = doc.core_properties.title or ""
            metadata["author"] = doc.core_properties.author or ""
    except (AttributeError, KeyError):
        pass

    # Extract text from paragraphs
    text_parts: list[str] = []
    text_parts.extend(para.text for para in doc.paragraphs if para.text.strip())

    # Also extract text from tables
    for table in doc.tables:
        table_rows = []
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                table_rows.append(row_text)
        if table_rows:
            text_parts.append("\n[Table]\n" + "\n".join(table_rows))

    return "\n\n".join(text_parts), metadata


def extract_text_from_xlsx(file_path: Path, max_rows: int = 1000) -> tuple[str, dict[str, Any]]:
    """
    Extract text summary from an Excel file using openpyxl.

    For data analysis, the agent should use execute_code with pandas instead.
    This provides a text overview of the spreadsheet structure.

    Args:
        file_path: Path to XLSX file
        max_rows: Maximum rows to show per sheet

    Returns:
        Tuple of (extracted_text, metadata_dict)
    """
    wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
    metadata = {
        "format": "xlsx",
        "sheets": wb.sheetnames,
    }

    text_parts = []

    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        text_parts.append(f"--- Sheet: {sheet_name} ---")

        rows_extracted = 0
        for row in sheet.iter_rows(values_only=True):
            if rows_extracted >= max_rows:
                text_parts.append(f"[... Truncated at {max_rows} rows ...]")
                break

            # Convert row to string, handling None values
            row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
            if row_text.strip():
                text_parts.append(row_text)
                rows_extracted += 1

    wb.close()
    return "\n".join(text_parts), metadata


def read_document(file_path: Path, max_chars: int = 100000) -> str:
    """
    Read a document and extract its text content.

    Handles PDF, DOCX, XLSX, and falls back to plain text for other formats.

    Args:
        file_path: Path to the document
        max_chars: Maximum characters to return (to avoid context bloat)

    Returns:
        Extracted text content with metadata header
    """
    file_path = Path(file_path)

    if not file_path.exists():
        return f"Error: File not found: {file_path}"

    suffix = file_path.suffix.lower()
    file_size_kb = file_path.stat().st_size / 1024

    try:
        if suffix == ".pdf":
            text, metadata = extract_text_from_pdf(file_path)
            header = f"[PDF: {file_path.name} | {metadata['pages']} pages | {file_size_kb:.1f} KB]"

        elif suffix in (".docx", ".doc"):
            if suffix == ".doc":
                return (
                    f"Error: Old .doc format not supported. "
                    f"Please convert {file_path.name} to .docx format."
                )
            text, metadata = extract_text_from_docx(file_path)
            header = f"[DOCX: {file_path.name} | {metadata['paragraphs']} paragraphs | {file_size_kb:.1f} KB]"

        elif suffix in (".xlsx", ".xls"):
            if suffix == ".xls":
                return (
                    f"Error: Old .xls format not supported. "
                    f"Please convert {file_path.name} to .xlsx format, "
                    f"or use execute_code with pandas to read it."
                )
            text, metadata = extract_text_from_xlsx(file_path)
            header = f"[XLSX: {file_path.name} | Sheets: {', '.join(metadata['sheets'])} | {file_size_kb:.1f} KB]"

        else:
            # Try to read as plain text
            try:
                with open(file_path, encoding="utf-8") as f:
                    text = f.read()
                header = f"[TEXT: {file_path.name} | {file_size_kb:.1f} KB]"
            except UnicodeDecodeError:
                return (
                    f"Error: Cannot read {file_path.name} as text. "
                    f"This appears to be a binary file with extension '{suffix}'. "
                    f"Supported document formats: PDF, DOCX, XLSX."
                )

        # Truncate if too long
        if len(text) > max_chars:
            text = text[:max_chars] + f"\n\n[... Truncated at {max_chars} characters ...]"

        return f"{header}\n\n{text}"

    except (OSError, ValueError, ImportError, RuntimeError) as e:
        logger.exception("Error reading document %s", file_path)
        return f"Error reading {file_path.name}: {type(e).__name__}: {e}"
