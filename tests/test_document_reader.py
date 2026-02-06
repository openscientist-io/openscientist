"""Tests for document_reader module."""

from pathlib import Path

import pytest
from shandy.document_reader import (
    extract_text_from_docx,
    extract_text_from_pdf,
    extract_text_from_xlsx,
    is_binary_document,
    is_text_file,
    read_document,
)


class TestFileTypeDetection:
    """Tests for file type detection functions."""

    def test_is_binary_document_pdf(self):
        assert is_binary_document(Path("test.pdf")) is True
        assert is_binary_document(Path("test.PDF")) is True

    def test_is_binary_document_docx(self):
        assert is_binary_document(Path("test.docx")) is True
        assert is_binary_document(Path("test.doc")) is True

    def test_is_binary_document_xlsx(self):
        assert is_binary_document(Path("test.xlsx")) is True
        assert is_binary_document(Path("test.xls")) is True

    def test_is_binary_document_text_files(self):
        assert is_binary_document(Path("test.csv")) is False
        assert is_binary_document(Path("test.txt")) is False
        assert is_binary_document(Path("test.json")) is False

    def test_is_text_file(self):
        assert is_text_file(Path("test.csv")) is True
        assert is_text_file(Path("test.txt")) is True
        assert is_text_file(Path("test.json")) is True
        assert is_text_file(Path("test.md")) is True
        assert is_text_file(Path("test.pdf")) is False


class TestPDFExtraction:
    """Tests for PDF text extraction."""

    @pytest.fixture
    def sample_pdf(self, tmp_path):
        """Create a sample PDF file for testing."""
        import fitz  # PyMuPDF

        pdf_path = tmp_path / "sample.pdf"
        doc = fitz.open()

        # Add a page with text
        page = doc.new_page()
        text = "This is a test PDF document.\nIt has multiple lines.\nAnd some scientific content."
        page.insert_text((50, 50), text)

        # Add a second page
        page2 = doc.new_page()
        page2.insert_text((50, 50), "This is page 2 of the document.")

        doc.save(pdf_path)
        doc.close()
        return pdf_path

    def test_extract_text_from_pdf(self, sample_pdf):
        text, metadata = extract_text_from_pdf(sample_pdf)

        assert "This is a test PDF document" in text
        assert "page 2" in text.lower()
        assert metadata["format"] == "pdf"
        assert metadata["pages"] == 2

    def test_extract_text_from_pdf_max_pages(self, sample_pdf):
        text, metadata = extract_text_from_pdf(sample_pdf, max_pages=1)

        assert "This is a test PDF document" in text
        assert "page 2" not in text.lower() or "Truncated" in text
        assert metadata["pages"] == 2  # Total pages, not extracted


class TestDocxExtraction:
    """Tests for DOCX text extraction."""

    @pytest.fixture
    def sample_docx(self, tmp_path):
        """Create a sample DOCX file for testing."""
        from docx import Document

        docx_path = tmp_path / "sample.docx"
        doc = Document()
        doc.add_paragraph("This is a test Word document.")
        doc.add_paragraph("It has multiple paragraphs.")
        doc.add_paragraph("With some scientific content about hypothesis testing.")

        # Add a simple table
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Header 1"
        table.cell(0, 1).text = "Header 2"
        table.cell(1, 0).text = "Value 1"
        table.cell(1, 1).text = "Value 2"

        doc.save(docx_path)
        return docx_path

    def test_extract_text_from_docx(self, sample_docx):
        text, metadata = extract_text_from_docx(sample_docx)

        assert "This is a test Word document" in text
        assert "hypothesis testing" in text
        assert metadata["format"] == "docx"
        assert metadata["paragraphs"] >= 3

    def test_extract_text_from_docx_includes_tables(self, sample_docx):
        text, metadata = extract_text_from_docx(sample_docx)

        assert "Header 1" in text
        assert "Value 1" in text


class TestXlsxExtraction:
    """Tests for XLSX text extraction."""

    @pytest.fixture
    def sample_xlsx(self, tmp_path):
        """Create a sample XLSX file for testing."""
        import openpyxl

        xlsx_path = tmp_path / "sample.xlsx"
        wb = openpyxl.Workbook()

        # First sheet (default)
        ws1 = wb.active
        ws1.title = "Data"
        ws1["A1"] = "Name"
        ws1["B1"] = "Value"
        ws1["A2"] = "Sample1"
        ws1["B2"] = 42.5
        ws1["A3"] = "Sample2"
        ws1["B3"] = 38.2

        # Second sheet
        ws2 = wb.create_sheet("Summary")
        ws2["A1"] = "Total"
        ws2["B1"] = 80.7

        wb.save(xlsx_path)
        return xlsx_path

    def test_extract_text_from_xlsx(self, sample_xlsx):
        text, metadata = extract_text_from_xlsx(sample_xlsx)

        assert "Name" in text
        assert "Sample1" in text
        assert "42.5" in text
        assert metadata["format"] == "xlsx"
        assert "Data" in metadata["sheets"]
        assert "Summary" in metadata["sheets"]

    def test_extract_text_from_xlsx_multiple_sheets(self, sample_xlsx):
        text, metadata = extract_text_from_xlsx(sample_xlsx)

        assert "Sheet: Data" in text
        assert "Sheet: Summary" in text
        assert "Total" in text


class TestReadDocument:
    """Tests for the main read_document function."""

    @pytest.fixture
    def sample_pdf(self, tmp_path):
        """Create a sample PDF file for testing."""
        import fitz

        pdf_path = tmp_path / "test.pdf"
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((50, 50), "PDF content here")
        doc.save(pdf_path)
        doc.close()
        return pdf_path

    @pytest.fixture
    def sample_txt(self, tmp_path):
        """Create a sample text file for testing."""
        txt_path = tmp_path / "test.txt"
        txt_path.write_text("Plain text content here")
        return txt_path

    def test_read_document_pdf(self, sample_pdf):
        result = read_document(sample_pdf)

        assert "[PDF:" in result
        assert "PDF content here" in result

    def test_read_document_text_file(self, sample_txt):
        result = read_document(sample_txt)

        assert "[TEXT:" in result
        assert "Plain text content here" in result

    def test_read_document_not_found(self, tmp_path):
        result = read_document(tmp_path / "nonexistent.pdf")

        assert "Error: File not found" in result

    def test_read_document_unsupported_binary(self, tmp_path):
        # Create a file with binary content but unknown extension
        binary_path = tmp_path / "test.bin"
        binary_path.write_bytes(b"\x00\x01\x02\x03\xff\xfe")

        result = read_document(binary_path)

        assert "Error" in result or "binary" in result.lower()

    def test_read_document_old_doc_format(self, tmp_path):
        # Create a fake .doc file
        doc_path = tmp_path / "test.doc"
        doc_path.write_bytes(b"fake doc content")

        result = read_document(doc_path)

        assert "Error" in result
        assert ".doc format not supported" in result

    def test_read_document_truncation(self, tmp_path):
        """Test that very long documents are truncated."""
        import fitz

        pdf_path = tmp_path / "long.pdf"
        doc = fitz.open()

        # Create a document with lots of text
        for i in range(50):
            page = doc.new_page()
            page.insert_text((50, 50), f"Page {i} " + "x" * 5000)

        doc.save(pdf_path)
        doc.close()

        result = read_document(pdf_path, max_chars=1000)

        assert "Truncated" in result
        assert len(result) < 2000  # Should be truncated


class TestEdgeCases:
    """Tests for edge cases and error handling."""

    def test_empty_pdf(self, tmp_path):
        """Test handling of empty PDF."""
        import fitz

        pdf_path = tmp_path / "empty.pdf"
        doc = fitz.open()
        doc.new_page()  # Empty page
        doc.save(pdf_path)
        doc.close()

        text, metadata = extract_text_from_pdf(pdf_path)

        assert metadata["pages"] == 1
        # Text might be empty or minimal

    def test_empty_docx(self, tmp_path):
        """Test handling of empty DOCX."""
        from docx import Document

        docx_path = tmp_path / "empty.docx"
        doc = Document()
        doc.save(docx_path)

        text, metadata = extract_text_from_docx(docx_path)

        assert metadata["format"] == "docx"
